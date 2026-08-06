"""Build the revised research report and executive brief."""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parents[1]
SENTIMENT_OUT = BASE / "outputs" / "sentiment"
OUT = BASE / "outputs"
OUT.mkdir(parents=True, exist_ok=True)

REPORT_DOCX = OUT / "Adoption_Without_Confidence_Research_Report_2026.docx"
BRIEF_DOCX = (
    OUT / "Favorable_Stance_Is_Not_Accuracy_Trust_Executive_Brief_2026.docx"
)

NAVY = RGBColor(16, 43, 51)
GREEN = RGBColor(35, 83, 71)
TEAL = RGBColor(46, 119, 111)
ORANGE = RGBColor(231, 111, 81)
GOLD = RGBColor(193, 145, 46)
INK = RGBColor(31, 31, 31)
GRAY = RGBColor(90, 98, 101)
LIGHT = "EEF4F1"
LIGHT_ORANGE = "FCEFEA"
TABLE_INDENT_DXA = 120


def set_font(run, size=None, color=INK, bold=None, italic=None):
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:ascii"), "Calibri")
    r_pr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_border(paragraph, color="2E776F", size="14"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def cell_margins(cell, value=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def configure_document(doc, business=False):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6 if business else 8)
    normal.paragraph_format.line_spacing = 1.10 if business else 1.30
    normal.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.LEFT if business else WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    tokens = {
        "Heading 1": (16, 17, 8),
        "Heading 2": (13, 11, 5),
        "Heading 3": (11.5, 8, 4),
    }
    for name, (size, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = GREEN if name != "Heading 3" else NAVY
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.38)
    bullet.paragraph_format.first_line_indent = Inches(-0.19)
    bullet.paragraph_format.space_after = Pt(4)
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.82)
        section.bottom_margin = Inches(0.78)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)


def page_furniture(doc, label):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(label)
        set_font(r, size=8.5, color=GRAY, bold=True)
        paragraph_border(p, color="D8E0DD", size="6")
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("Imagination Applied  |  2026  |  ")
        set_font(r, size=8.5, color=GRAY)
        add_page_field(p)


def body(doc, text, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, italic=italic)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)


def callout(doc, label, text, fill=LIGHT):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(9)
    shade_paragraph(p, fill)
    r = p.add_run(label.upper() + "  ")
    set_font(r, size=10.5, color=GREEN, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p


def add_table(doc, headers, rows, widths, numeric=(), font_size=8.8):
    numeric = set(numeric)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    repeat_header(header)
    for index, label in enumerate(headers):
        cell = header.cells[index]
        shade_cell(cell, "E4ECE9")
        p = cell.paragraphs[0]
        p.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
            if index in numeric
            else WD_ALIGN_PARAGRAPH.LEFT
        )
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(label))
        set_font(r, size=font_size, color=NAVY, bold=True)
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            p = row.cells[index].paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if index in numeric
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_font(r, size=font_size)
    table_geometry(table, widths)
    return table


def source_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=8.4, color=GRAY, italic=True)


def picture(doc, path, width, caption, alt_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", caption)
    p.paragraph_format.space_after = Pt(3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(caption)
    set_font(r, size=8.6, color=GRAY, italic=True)


def hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E776F")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    link.append(run)
    paragraph._p.append(link)


def reference(doc, text, url):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text + " ")
    set_font(r, size=9.3)
    hyperlink(p, url, url)


def cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("IMAGINATION APPLIED RESEARCH SERIES")
    set_font(r, size=10, color=GOLD, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Adoption Without Confidence?")
    set_font(r, size=30, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run(
        "Favorable Stance, Accuracy Trust, and Reported AI-Use Frequency "
        "in the 2025 Stack Overflow Survey"
    )
    set_font(r, size=14.5, color=TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(45)
    r = p.add_run("Josh Penzell  |  Imagination Applied")
    set_font(r, size=11.5, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("July 2026")
    set_font(r, size=10.5, color=GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(65)
    paragraph_border(p, size="18")
    r = p.add_run(
        "A denominator-aware secondary analysis of the 2023-2025 "
        "Stack Overflow Developer Surveys"
    )
    set_font(r, size=9.5, color=GRAY, italic=True)


def model_row(models, name):
    return models[models["model"] == name].iloc[0]


def load_results():
    return {
        "current_rates": pd.read_csv(
            SENTIMENT_OUT / "current_user_daily_use_by_stance.csv"
        ),
        "full_rates": pd.read_csv(
            SENTIMENT_OUT / "full_sample_daily_use_by_stance.csv"
        ),
        "high_distrust_rates": pd.read_csv(
            SENTIMENT_OUT
            / "current_user_high_distrust_daily_use_by_stance.csv"
        ),
        "primary_models": pd.read_csv(
            SENTIMENT_OUT / "primary_current_user_model_comparison.csv"
        ),
        "full_models": pd.read_csv(
            SENTIMENT_OUT / "full_sample_model_comparison.csv"
        ),
        "professional_models": pd.read_csv(
            SENTIMENT_OUT
            / "professional_current_user_model_comparison.csv"
        ),
        "repeated": pd.read_csv(
            SENTIMENT_OUT / "primary_repeated_cv_comparisons.csv"
        ),
        "country": pd.read_csv(
            SENTIMENT_OUT / "primary_country_grouped_cv.csv"
        ),
        "heldout": json.loads(
            (SENTIMENT_OUT / "primary_heldout_metrics.json").read_text(
                encoding="utf-8"
            )
        ),
    }


def add_rate_table(doc, rates, caption):
    rows = [
        (
            row.favorable_stance,
            f"{int(row.n):,}",
            f"{row.daily_rate * 100:.1f}%",
            f"{row.ci_low * 100:.1f}-{row.ci_high * 100:.1f}%",
        )
        for row in rates.itertuples()
    ]
    add_table(
        doc,
        ["Favorable stance", "n", "Daily use", "95% Wilson interval"],
        rows,
        [3650, 1400, 1700, 2610],
        numeric=(1, 2, 3),
    )
    source_note(doc, caption)


def build_report():
    r = load_results()
    current_rates = r["current_rates"]
    full_rates = r["full_rates"]
    high_distrust = r["high_distrust_rates"]
    primary = r["primary_models"]
    full = r["full_models"]
    professional = r["professional_models"]
    repeated = r["repeated"]
    country = r["country"]
    heldout = r["heldout"]

    stance = model_row(primary, "Stance only")
    trust = model_row(primary, "Trust only")
    context_trust = model_row(primary, "Context + trust")
    context_stance = model_row(primary, "Context + stance")
    complete = model_row(primary, "Context + trust + stance")

    report = Document()
    configure_document(report)
    page_furniture(report, "ADOPTION WITHOUT CONFIDENCE?  |  RESEARCH REPORT")
    cover(report)
    report.add_page_break()

    report.add_heading("Abstract", level=1)
    body(
        report,
        "Stack Overflow has already reported that AI use rose while trust weakened. "
        "This secondary analysis asks a narrower question: when a broad favorable "
        "stance toward using AI and trust in AI-output accuracy are measured "
        "separately, which carries more information about reported use frequency? "
        "The three-year context uses 203,812 CSV records from repeated cross-sections. "
        "The 2025 complete-case sample contains 33,231 respondents, of whom 25,698 "
        "(77.3%) identify as professional developers. To reduce the direct contrast "
        "between users and nonusers, the primary analysis is restricted to 26,102 "
        "current AI users and classifies daily versus weekly or less-frequent use. "
        "Daily use ranges from 18.8% among very unfavorable current users to 88.1% "
        "among very favorable current users, although the middle categories are not "
        "strictly ordered. A categorical stance-only model reaches five-fold "
        "cross-validated ROC AUC 0.758, compared with 0.669 for accuracy trust alone. "
        "Context plus trust reaches AUC 0.704; context plus stance reaches 0.790; "
        "adding trust raises that result to 0.793. Across 50 matched repeated splits, "
        "context plus stance exceeds context plus trust by 0.087 AUC on average, "
        "while trust adds 0.003 after stance. The full-sample gradient is larger "
        "(3.4% to 85.7%; stance AUC 0.822), consistent with construct proximity "
        "between favoring AI use and reporting AI use. These same-wave associations "
        "do not establish causal direction or future predictive validity.",
    )
    callout(
        report,
        "Central finding",
        "Among people already using AI, favorable stance carries substantially more "
        "information about daily versus less-frequent use than accuracy trust does. "
        "That shows the two survey items are not interchangeable here; it does not "
        "show that changing stance will change behavior.",
    )
    report.add_heading("Key findings", level=2)
    bullets(
        report,
        [
            "Among 26,102 current users, daily use is 18.8% for the very unfavorable and 88.1% for the very favorable; unsure (34.0%) is slightly above indifferent (31.5%), so the pattern is steep but not perfectly monotonic.",
            "Stance alone classifies daily versus less-frequent use with AUC 0.758, compared with 0.669 for accuracy trust alone.",
            "Context plus stance outperforms context plus trust in all 50 matched repeated splits; adding trust after stance produces a small but consistently positive improvement.",
            "Among current users who highly distrust AI accuracy, daily use ranges from 17.7% for the very unfavorable to 85.1% for the very favorable.",
            "The result remains similar among professional current users and in five-fold country-grouped validation, but these are internal robustness checks rather than external validation.",
            "The strongest counterinterpretation remains viable: the stance item is broader and closer to the use outcome than the narrower accuracy-trust item, and use may shape stance as much as stance shapes use.",
        ],
    )

    report.add_heading("1. What is new - and what is not", level=1)
    body(
        report,
        "The broad adoption-with-low-trust story is not new. Stack Overflow has "
        "published the 2025 use, favorable-stance, and accuracy-trust distributions, "
        "related favorable stance to workflow integration, and later discussed the "
        "developer AI trust gap directly. Broader workplace research has also "
        "distinguished AI use, attitudes, and trust.",
    )
    body(
        report,
        "The contribution here is narrower: a reproducible large-sample comparison "
        "of how the Stack Overflow favorable-stance and accuracy-trust items classify "
        "reported use frequency, with the main test restricted to current users and "
        "supported by trust-stratified, professional-status, alternative-outcome, "
        "country-grouped, calibration, and repeated-split checks. The analysis does "
        "not claim that stance and trust are newly discovered constructs.",
    )
    report.add_heading("Research questions", level=2)
    bullets(
        report,
        [
            "RQ1. What remains of the 2023-2025 use-trust pattern after harmonizing the trust denominator?",
            "RQ2. Among current AI users, how does reported daily use vary across favorable-stance responses?",
            "RQ3. How much classification information does favorable stance add beyond accuracy trust and respondent context?",
            "RQ4. Does the result remain visible among professional developers, within trust levels, and when countries form held-out groups?",
            "RQ5. Which managerial implications survive the causal and measurement limitations?",
        ],
    )

    report.add_heading("2. Construct boundaries and closest prior work", level=1)
    report.add_heading("Favorable stance is not a validated sentiment scale", level=2)
    body(
        report,
        "The Stack Overflow item AISent asks how favorable the respondent's stance "
        "is toward using AI tools in the development workflow. Stack Overflow labels "
        "the item AI tool sentiment. This report uses 'favorable stance' for analytic "
        "claims because the item is a single broad evaluation, not a validated "
        "multi-item affect scale. It may combine usefulness, fit, identity, "
        "opportunity, prior experience, and willingness to use.",
    )
    report.add_heading("Accuracy trust is narrower", level=2)
    body(
        report,
        "The AIAcc item asks how much the respondent trusts the accuracy of AI "
        "output. Trust research distinguishes attitude, reliance, and calibrated "
        "reliance. A person can find AI useful enough to use daily while distrusting "
        "individual answers enough to test them. Confidence in accuracy also does "
        "not guarantee that a tool is useful, available, permitted, or relevant.",
    )
    report.add_heading("Closest empirical comparison", level=2)
    body(
        report,
        "Choudhuri and colleagues used validated constructs and a theoretically "
        "grounded model with 238 developers at GitHub and Microsoft. They found that "
        "system/output quality, functional value, and goal maintenance shaped trust, "
        "and that trust plus cognitive styles related to intentions to use GenAI. "
        "That work is theoretically and psychometrically stronger. The present "
        "analysis is larger, public, and focused on direct comparative classification "
        "of a reported frequency category. The approaches are complementary.",
    )

    report.add_heading("3. Data and methods", level=1)
    body(
        report,
        "The public 2023, 2024, and 2025 Stack Overflow Developer Survey CSVs "
        "contain 203,812 rows: 89,184 in 2023, 65,437 in 2024, and 49,191 in the "
        "downloaded 2025 file. Stack Overflow's published 2025 totals differ "
        "slightly from the public CSV, consistent with versioning or post-release "
        "filtering. Full hashes and row counts are documented in the repository.",
    )
    body(
        report,
        "These are voluntary nonprobability samples. 'Stack Overflow survey "
        "respondents' is the population label used throughout. The sample includes "
        "professional developers, learners, hobbyists, former developers, and other "
        "people who code. Results are not population prevalence estimates.",
    )
    sample_rows = [
        ("2025 total CSV rows", "49,191", "Survey file"),
        ("Answered AI-use item", "33,720", "68.5% of CSV rows"),
        ("Complete use, stance, and trust", "33,231", "Full-sample sensitivity"),
        ("Professional developers", "25,698", "77.3% of complete cases"),
        ("Current AI users", "26,102", "Primary analysis"),
        ("Professional current users", "20,760", "Population-label sensitivity"),
    ]
    add_table(
        report,
        ["Sample", "n", "Role in analysis"],
        sample_rows,
        [4200, 1500, 3660],
        numeric=(1,),
    )
    source_note(
        report,
        "Table 1. Analysis samples from the downloaded 2025 public CSV.",
    )
    report.add_heading("Measures and primary outcome", level=2)
    bullets(
        report,
        [
            "Primary outcome: among current users, daily use versus weekly, monthly, or infrequent use.",
            "Full-sample sensitivity: daily use versus every other answered AI-use response, including nonuse.",
            "Favorable stance: the six AISent labels are modeled categorically; no equal spacing is assumed.",
            "Accuracy trust: the five AIAcc labels are modeled categorically.",
            "Context: work and coding experience, age, professional status, employment group, role, organization size, work mode, perceived AI threat, and country. Infrequent-country handling is learned inside each training fold.",
        ],
    )
    callout(
        report,
        "Deliberate exclusion",
        "AI-agent use, workflow integration, task-complexity evaluations, "
        "frustration items, and learning-route fields are excluded because they are "
        "conditional on, downstream from, or behaviorally too close to current use.",
        fill=LIGHT_ORANGE,
    )
    report.add_heading("Analytic strategy", level=2)
    body(
        report,
        "Six logistic specifications use common five-fold stratified splits: stance "
        "only; trust only; context; context plus trust; context plus stance; and "
        "context plus trust plus stance. Metrics include accuracy, balanced accuracy, "
        "ROC AUC, log loss, and Brier score. Preprocessing occurs within each fold.",
    )
    body(
        report,
        "Robustness checks include 50 matched repeated splits, a separate 80/20 "
        "held-out calibration and permutation check, country-grouped folds, "
        "professional-current-user restriction, alternative outcome thresholds, and "
        "the full complete-case sample. None identifies causal direction.",
    )

    report.add_heading("4. Denominator-aware three-year context", level=1)
    body(
        report,
        "Among respondents who answered the AI-use item, current use increased from "
        "44.4% in 2023 to 61.8% in 2024 and 78.5% in 2025, while favorable stance "
        "fell from 76.3% to 59.7%. The 2025 use item had a lower response rate and "
        "introduced frequency categories. These are repeated cross-sections, not a "
        "panel or a clean population trend.",
    )
    trend_rows = [
        ("Current use among AI-use respondents", "44.4%", "61.8%", "78.5%", "Wording and response-rate caveat"),
        ("Favorable stance", "76.3%", "72.0%", "59.7%", "Repeated cross-sections"),
        ("Trust among current users", "48.2%", "43.0%", "39.3%", "Harmonized denominator"),
        ("Distrust among current users", "24.2%", "30.4%", "37.3%", "Harmonized denominator"),
    ]
    add_table(
        report,
        ["Indicator", "2023", "2024", "2025", "Interpretation"],
        trend_rows,
        [3300, 1050, 1050, 1050, 2910],
        numeric=(1, 2, 3),
    )
    source_note(
        report,
        "Table 2. Repeated cross-sectional context. Trust is restricted to current users for comparability.",
    )
    picture(
        report,
        SENTIMENT_OUT / "three_year_context.png",
        6.25,
        "Figure 1. Reported use rose while favorability and harmonized current-user trust weakened.",
        "Line chart for 2023 through 2025 showing current AI use rising, favorable stance falling, and trust among current users weakening modestly.",
    )
    body(
        report,
        "The trust routing change matters. In 2023 and 2024, the trust item was "
        "routed to current users. In 2025, nonusers also answered it. Restricting "
        "2025 to current users yields 39.3% trust and 37.3% distrust, compared with "
        "43.0% and 30.4% in 2024.",
    )

    report.add_heading("5. Primary result: frequency among current users", level=1)
    body(
        report,
        "The primary restriction removes nonusers and people who only plan to use "
        "AI. Within the 26,102 current users, daily use is 18.8% for a very "
        "unfavorable stance and 88.1% for a very favorable stance. The middle "
        "categories are not strictly monotonic: unsure is 34.0% and indifferent is "
        "31.5%. The defensible description is steep separation between unfavorable "
        "and favorable positions, not a perfectly even ordinal ladder.",
    )
    add_rate_table(
        report,
        current_rates,
        "Table 3. Daily versus less-frequent use among current AI users, n=26,102.",
    )
    picture(
        report,
        SENTIMENT_OUT / "primary_current_user_stance_gradient.png",
        6.25,
        "Figure 2. Observed and context-adjusted daily-use probabilities among current AI users.",
        "Line chart showing large differences in daily use across favorable-stance categories among current AI users, with small reversals in the middle categories.",
    )

    report.add_heading("6. Comparative classification", level=1)
    body(
        report,
        f"Among current users, stance alone reaches AUC {stance.roc_auc_mean:.3f}; "
        f"trust alone reaches {trust.roc_auc_mean:.3f}. Context plus trust reaches "
        f"{context_trust.roc_auc_mean:.3f}, while context plus stance reaches "
        f"{context_stance.roc_auc_mean:.3f}. The complete model reaches "
        f"{complete.roc_auc_mean:.3f}. Accuracy trust is not irrelevant, but it adds "
        "little frequency-classification information after stance is known.",
    )
    order = [
        "Majority baseline",
        "Stance only",
        "Trust only",
        "Context",
        "Context + trust",
        "Context + stance",
        "Context + trust + stance",
    ]
    rows = []
    for name in order:
        row = model_row(primary, name)
        rows.append(
            (
                name,
                f"{row.accuracy_mean:.3f}",
                f"{row.roc_auc_mean:.3f}",
                f"{row.log_loss_mean:.3f}",
                f"{row.brier_mean:.3f}",
            )
        )
    add_table(
        report,
        ["Model", "Accuracy", "ROC AUC", "Log loss", "Brier"],
        rows,
        [4050, 1300, 1300, 1350, 1360],
        numeric=(1, 2, 3, 4),
    )
    source_note(
        report,
        "Table 4. Five-fold cross-validation among current users. Higher AUC and accuracy are better; lower log loss and Brier are better.",
    )
    picture(
        report,
        SENTIMENT_OUT / "primary_current_user_model_auc.png",
        6.25,
        "Figure 3. Cross-validated AUC for daily versus less-frequent use among current users.",
        "Horizontal bar chart showing that models containing favorable stance outperform models containing accuracy trust or respondent context without stance.",
    )
    repeat_stance = repeated.iloc[0]
    repeat_trust = repeated.iloc[1]
    body(
        report,
        f"Across 50 matched repeated splits, context plus stance exceeds context "
        f"plus trust by {repeat_stance.auc_delta_mean:.3f} AUC on average "
        f"(2.5th to 97.5th empirical split percentiles: "
        f"{repeat_stance.auc_delta_2_5pct:.3f} to "
        f"{repeat_stance.auc_delta_97_5pct:.3f}) and wins in all 50 splits. "
        f"Adding trust after stance improves AUC by "
        f"{repeat_trust.auc_delta_mean:.3f} on average and is positive in all 50 "
        "splits. These overlapping resamples are stability summaries, not "
        "independent confidence intervals.",
    )
    body(
        report,
        f"On a separate 20% holdout, the complete model reaches AUC "
        f"{heldout['roc_auc']:.3f}, calibration slope "
        f"{heldout['calibration_slope']:.2f}, and mean absolute observed-versus-"
        f"predicted decile gap {heldout['mean_absolute_decile_gap']:.3f}. This "
        "supports internal calibration, not external transportability.",
    )

    report.add_heading("7. Robustness and full-sample amplification", level=1)
    full_stance = model_row(full, "Stance only")
    full_trust = model_row(full, "Trust only")
    pro_stance = model_row(professional, "Stance only")
    pro_trust = model_row(professional, "Trust only")
    country_stance = country[country["model"] == "Stance only"].iloc[0]
    country_trust = country[country["model"] == "Trust only"].iloc[0]
    low_full = full_rates.iloc[0]
    high_full = full_rates.iloc[-1]
    body(
        report,
        f"In the full 33,231-person complete-case sample, daily use runs from "
        f"{low_full.daily_rate * 100:.1f}% to {high_full.daily_rate * 100:.1f}% "
        f"across the stance extremes, and stance AUC is "
        f"{full_stance.roc_auc_mean:.3f} versus {full_trust.roc_auc_mean:.3f} for "
        "trust. This is valid descriptively, but the larger separation partly "
        "reflects mixing users and nonusers while the stance item asks about "
        "favoring AI use.",
    )
    robustness_rows = [
        ("Current users only", "26,102", f"{stance.roc_auc_mean:.3f}", f"{trust.roc_auc_mean:.3f}", "Primary frequency comparison"),
        ("Professional current users", "20,760", f"{pro_stance.roc_auc_mean:.3f}", f"{pro_trust.roc_auc_mean:.3f}", "Population-label sensitivity"),
        ("Country-grouped current users", f"{int(country_stance.n):,}", f"{country_stance.grouped_auc_mean:.3f}", f"{country_trust.grouped_auc_mean:.3f}", f"{int(country_stance.countries)} countries"),
        ("Full complete-case sample", "33,231", f"{full_stance.roc_auc_mean:.3f}", f"{full_trust.roc_auc_mean:.3f}", "Includes nonusers; amplified"),
    ]
    add_table(
        report,
        ["Check", "n", "Stance AUC", "Trust AUC", "Interpretation"],
        robustness_rows,
        [3000, 1150, 1250, 1250, 2710],
        numeric=(1, 2, 3),
        font_size=8.6,
    )
    source_note(
        report,
        "Table 5. Robustness checks. Country-grouped folds are not external organizational validation.",
    )
    body(
        report,
        "Stance classification is stronger for broader thresholds such as any current "
        "use or current use plus plans. That consistency is also a warning: the "
        "broader the target becomes adoption versus nonadoption, the closer it is to "
        "the wording of a favorable stance toward using AI.",
    )

    report.add_heading("8. Favorable stance is not accuracy trust", level=1)
    distrust_low = high_distrust.iloc[0]
    distrust_high = high_distrust.iloc[-1]
    body(
        report,
        "The within-trust comparison is now restricted to current users. Among "
        "current users who highly distrust AI accuracy, daily use is "
        f"{distrust_low.daily_rate * 100:.1f}% for the very unfavorable "
        f"(n={int(distrust_low.n):,}) and "
        f"{distrust_high.daily_rate * 100:.1f}% for the very favorable "
        f"(n={int(distrust_high.n):,}). Broad favorability can coexist with "
        "skepticism about individual outputs.",
    )
    picture(
        report,
        SENTIMENT_OUT / "primary_current_user_stance_trust_heatmap.png",
        6.35,
        "Figure 4. Daily-use rates across favorable stance and accuracy trust among current AI users.",
        "Heat map showing daily AI use across six stance levels and five accuracy-trust levels among current AI users.",
    )
    callout(
        report,
        "Interpret carefully",
        "Restricting to current users removes the nonuser contrast but conditions on "
        "adoption. It is a tougher descriptive sensitivity check, not a causal fix.",
    )

    report.add_heading("9. The strongest counterarguments", level=1)
    report.add_heading("The stance item is broader and closer to the outcome", level=2)
    body(
        report,
        "A broad question about favoring AI use should align more closely with "
        "self-reported AI use than a narrow question about accuracy. The evidence "
        "supports 'accuracy trust alone is an incomplete summary of use orientation,' "
        "not 'stance is psychologically more important.'",
    )
    report.add_heading("Use may shape stance", level=2)
    body(
        report,
        "Daily use may make a respondent more favorable through experienced value or "
        "less favorable through repeated failures. Temporal order and reciprocal "
        "effects cannot be recovered from one survey wave.",
    )
    report.add_heading("Same-wave self-report can inflate coherence", level=2)
    body(
        report,
        "Stance, trust, and use are reported in the same session. Consistency motives, "
        "shared interpretation, and omitted variables can strengthen alignment. "
        "Cross-validation and calibration do not remove common-method bias.",
    )
    report.add_heading("Classification importance is not managerial importance", level=2)
    body(
        report,
        "Accuracy trust adds little AUC after stance, but it remains essential for "
        "verification, over-reliance, and safety. A variable can matter greatly to "
        "outcomes while adding little information about use frequency.",
    )

    report.add_heading("10. What organizations can responsibly do", level=1)
    body(
        report,
        "The results justify separate measurement and longitudinal testing, not "
        "persuasion campaigns. The management question is not how to maximize "
        "enthusiasm or trust. It is how to create useful, appropriately skeptical, "
        "well-verified use.",
    )
    dashboard_rows = [
        ("Behavior", "Frequency, task, workflow depth", "What people do"),
        ("Favorable stance", "Broad evaluation of using AI", "Perceived value, fit, or willingness"),
        ("Accuracy trust", "Confidence in output correctness", "Belief about claims and answers"),
        ("Verification burden", "Testing, review, rework, defects", "Cost of appropriate skepticism"),
        ("Outcomes", "Quality, speed, learning, risk", "Whether use creates value safely"),
    ]
    add_table(
        report,
        ["Dimension", "Example measures", "What it answers"],
        dashboard_rows,
        [1700, 3660, 4000],
    )
    source_note(
        report,
        "Table 6. An analytically distinct adoption-measurement dashboard.",
    )
    report.add_heading("Test interventions longitudinally", level=2)
    bullets(
        report,
        [
            "Measure stance, accuracy trust, expected value, and verification practices before a change.",
            "Use behavioral telemetry and quality outcomes alongside self-report.",
            "Randomize or stagger workflow support, task examples, or verification scaffolds where feasible.",
            "Re-measure after real work rather than immediately after a demonstration.",
            "Look for calibrated reliance rather than maximum trust.",
        ],
    )

    report.add_heading("11. Limitations", level=1)
    bullets(
        report,
        [
            "Voluntary nonprobability sample; prevalence does not represent all developers.",
            "Only 77.3% of complete cases identify as professional developers.",
            "Repeated cross-sections; no individual change is observed across years.",
            "Routing and wording changes; the 2025 AI-module response rate is lower.",
            "Single-item constructs; favorable stance and accuracy trust are not validated scales.",
            "Same-wave self-report; temporal order, common-method bias, and objective behavior are unresolved.",
            "Criterion proximity; favoring AI use is semantically closer to use frequency than accuracy trust is.",
            "Conditioning; the current-user restriction reduces one artifact but is not a causal design.",
            "Model transportability; internal and country-grouped resampling do not prove performance in a new organization or future year.",
            "Unmeasured confounding; availability, mandates, task mix, tool quality, usefulness, and prior experience may explain the association.",
            "No objective calibration measure; the survey cannot determine whether trust is warranted.",
        ],
    )

    report.add_heading("12. Conclusion", level=1)
    body(
        report,
        "The familiar story is adoption rising while trust falls. The more precise "
        "finding is that a broad favorable stance and accuracy-specific trust carry "
        "different information about use frequency. Even among people already using "
        "AI, favorable stance separates daily from less-frequent use substantially "
        "better than accuracy trust does. The result survives professional-current-"
        "user and country-grouped checks.",
    )
    body(
        report,
        "The same evidence blocks the tempting causal story. Favorable stance is "
        "broad, criterion-proximal, and potentially shaped by use. It should be "
        "treated as an analytically distinct survey item, not a lever "
        "proven to drive adoption. Organizations should measure stance, trust, "
        "behavior, verification, and outcomes separately, then test what changes.",
    )
    callout(
        report,
        "Bottom line",
        "Access is not adoption. Favoring use is not trusting output. Frequent use is "
        "not calibrated reliance. Measure each one - then test what changes.",
    )

    report.add_heading("Technical appendix", level=1)
    report.add_heading("A.1 Reproducibility", level=2)
    body(
        report,
        "The repository includes source-data hashes, recodes, denominator notes, "
        "machine-readable results, analysis and publication scripts, and figures. "
        "Raw survey files are excluded because of size and licensing; official "
        "download paths and integrity checks are documented.",
    )
    report.add_heading("A.2 Claim ladder", level=2)
    claim_rows = [
        ("Supported descriptively", "Daily-use rates differ sharply across stance categories among current users."),
        ("Supported within this survey", "Stance adds substantially more cross-validated frequency-classification information than accuracy trust."),
        ("Plausible but unresolved", "Stance influences use, use influences stance, or both."),
        ("Not supported", "Changing employee stance will cause greater or better adoption."),
        ("Not supported", "Trust is unimportant because it adds little AUC after stance."),
    ]
    add_table(
        report,
        ["Evidence level", "Claim"],
        claim_rows,
        [2700, 6660],
        font_size=9.0,
    )
    source_note(
        report,
        "Table A1. Boundaries for publication and public communication.",
    )

    report.add_heading("References", level=1)
    references = [
        (
            "Choudhuri, R., Trinkenreich, B., Pandita, R., Kalliamvakou, E., Steinmacher, I., Gerosa, M., Sanchez, C. A., & Sarma, A. (2025). What guides our choices? Modeling developers' trust and behavioral intentions towards GenAI. Proceedings of ICSE 2025, 1691-1703.",
            "https://doi.org/10.1109/ICSE55347.2025.00087",
        ),
        (
            "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319-340.",
            "https://doi.org/10.2307/249008",
        ),
        (
            "Hoff, K. A., & Bashir, M. (2015). Trust in automation: Integrating empirical evidence on factors that influence trust. Human Factors, 57(3), 407-434.",
            "https://doi.org/10.1177/0018720814547570",
        ),
        (
            "Lee, J. D., & See, K. A. (2004). Trust in automation: Designing for appropriate reliance. Human Factors, 46(1), 50-80.",
            "https://doi.org/10.1518/hfes.46.1.50_30392",
        ),
        (
            "Gillespie, N., Lockey, S., Ward, T., Macdade, A., & Hassed, G. (2025). Trust, attitudes and use of artificial intelligence: A global study 2025. University of Melbourne & KPMG.",
            "https://doi.org/10.26188/28822919",
        ),
        (
            "Stack Overflow. (2023). 2023 Developer Survey.",
            "https://survey.stackoverflow.co/2023",
        ),
        (
            "Stack Overflow. (2024). 2024 Developer Survey: AI.",
            "https://survey.stackoverflow.co/2024/ai",
        ),
        (
            "Stack Overflow. (2025). 2025 Developer Survey: AI.",
            "https://survey.stackoverflow.co/2025/ai",
        ),
        (
            "Stack Overflow. (2026, February 18). Mind the gap: Closing the AI trust gap for developers.",
            "https://stackoverflow.blog/2026/02/18/closing-the-developer-ai-trust-gap/",
        ),
        (
            "Venkatesh, V., Morris, M. G., Davis, G. B., & Davis, F. D. (2003). User acceptance of information technology: Toward a unified view. MIS Quarterly, 27(3), 425-478.",
            "https://doi.org/10.2307/30036540",
        ),
        (
            "Wang, R., Cheng, R., Ford, D., & Zimmermann, T. (2024). Investigating and designing for trust in AI-powered code generation tools. Proceedings of FAccT 2024.",
            "https://doi.org/10.1145/3630106.3658984",
        ),
    ]
    for text, url in references:
        reference(report, text, url)

    report.core_properties.title = "Adoption Without Confidence?"
    report.core_properties.subject = (
        "Favorable stance, accuracy trust, and reported AI-use frequency "
        "in the 2025 Stack Overflow survey"
    )
    report.core_properties.author = "Josh Penzell"
    report.core_properties.keywords = (
        "AI adoption, favorable stance, accuracy trust, Stack Overflow, "
        "reported use frequency, calibrated reliance"
    )
    report.core_properties.comments = (
        "Imagination Applied Research Series, revised July 2026"
    )
    report.save(REPORT_DOCX)


def build_brief():
    r = load_results()
    rates = r["current_rates"]
    models = r["primary_models"]
    high_distrust = r["high_distrust_rates"]
    stance = model_row(models, "Stance only")
    trust = model_row(models, "Trust only")
    context_trust = model_row(models, "Context + trust")
    context_stance = model_row(models, "Context + stance")
    complete = model_row(models, "Context + trust + stance")

    brief = Document()
    configure_document(brief, business=True)
    page_furniture(
        brief,
        "FAVORABLE STANCE IS NOT ACCURACY TRUST  |  EXECUTIVE BRIEF",
    )
    p = brief.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("EXECUTIVE BRIEF")
    set_font(r, size=10, color=GOLD, bold=True)
    p = brief.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Favorable Stance Is Not Accuracy Trust")
    set_font(r, size=27, color=NAVY, bold=True)
    p = brief.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(
        "What 26,102 current AI users in the 2025 Stack Overflow survey "
        "reveal about reported use frequency"
    )
    set_font(r, size=13.3, color=TEAL)
    paragraph_border(p)

    low = rates.iloc[0]
    high = rates.iloc[-1]
    callout(
        brief,
        "The finding",
        f"Among people already using AI, reported daily use is "
        f"{low.daily_rate * 100:.1f}% for the very unfavorable and "
        f"{high.daily_rate * 100:.1f}% for the very favorable. Favorable stance "
        "carries more classification information about daily versus less-frequent "
        "use than accuracy trust does. This is an association, not a causal effect.",
    )
    brief.add_heading("The numbers that matter", level=1)
    metrics = [
        (f"{stance.roc_auc_mean:.3f}", "Stance-only ROC AUC"),
        (f"{trust.roc_auc_mean:.3f}", "Trust-only ROC AUC"),
        (
            f"{context_trust.roc_auc_mean:.3f} -> {context_stance.roc_auc_mean:.3f}",
            "Context + trust versus context + stance",
        ),
        (f"{complete.roc_auc_mean:.3f}", "Complete-model ROC AUC"),
    ]
    add_table(
        brief,
        ["Result", "Meaning"],
        metrics,
        [2300, 7060],
        numeric=(0,),
        font_size=10.0,
    )
    source_note(
        brief,
        "Five-fold cross-validation among current users, n=26,102. Same-survey classification, not future forecasting.",
    )
    picture(
        brief,
        SENTIMENT_OUT / "primary_current_user_stance_gradient.png",
        6.0,
        "Daily-use rates and context-adjusted probabilities across favorable-stance responses among current users.",
        "Line chart showing large differences in daily AI use across favorable-stance responses among current users.",
    )
    brief.add_heading("The important distinction", level=1)
    body(
        brief,
        "Favorable stance is a broad evaluation of using AI. Accuracy trust is a "
        "narrow belief about whether AI output is correct. The broader, more "
        "use-proximal item aligns more closely with frequency. That does not make "
        "stance a proven psychological driver or make trust unimportant.",
    )
    distrust_low = high_distrust.iloc[0]
    distrust_high = high_distrust.iloc[-1]
    callout(
        brief,
        "A tougher comparison",
        f"Among current users who highly distrust AI accuracy, daily use is "
        f"{distrust_low.daily_rate * 100:.1f}% for the very unfavorable and "
        f"{distrust_high.daily_rate * 100:.1f}% for the very favorable. A user can "
        "value AI while remaining skeptical enough to verify its outputs.",
        fill=LIGHT_ORANGE,
    )
    brief.add_heading("What this changes for leaders", level=1)
    moves = [
        ("1", "Track behavior, not access.", "Measure frequency, task, workflow depth, and outcomes."),
        ("2", "Separate stance from accuracy trust.", "One reflects broad value and fit; the other confidence in correctness."),
        ("3", "Optimize for calibrated reliance.", "Seek useful, appropriately skeptical, well-verified use."),
        ("4", "Test changes longitudinally.", "Use pre/post measures, telemetry, quality, rework, and staggered support."),
    ]
    add_table(
        brief,
        ["", "Move", "Why"],
        moves,
        [600, 3500, 5260],
        font_size=9.1,
    )
    brief.add_heading("What the data do not prove", level=1)
    bullets(
        brief,
        [
            "That favorable stance causes daily use; use may shape stance.",
            "That the favorable-stance item is a validated sentiment scale.",
            "That trust is unimportant because it adds little classification information.",
            "That training or persuasion will increase adoption.",
            "That Stack Overflow respondents represent all developer workforces.",
        ],
    )
    brief.add_heading("A better adoption dashboard", level=1)
    dashboard = [
        ("Behavior", "How often, where, and for what"),
        ("Favorable stance", "Whether AI feels valuable and workable"),
        ("Accuracy trust", "Confidence that output is correct"),
        ("Verification burden", "Testing, review, rework, defects, and time"),
        ("Outcomes", "Quality, speed, learning, risk, and customer impact"),
    ]
    add_table(
        brief,
        ["Dimension", "Question"],
        dashboard,
        [2700, 6660],
        font_size=9.5,
    )
    callout(
        brief,
        "Bottom line",
        "Favoring AI use is not the same as trusting AI output. Measure both - "
        "alongside behavior, verification, and outcomes - then test what changes.",
    )
    p = brief.add_paragraph()
    r = p.add_run(
        "Source: Penzell, J. (2026). Adoption Without Confidence? "
        "Imagination Applied Research Series. "
    )
    set_font(r, size=8.5, color=GRAY)
    hyperlink(
        p,
        "Stack Overflow 2025 AI results",
        "https://survey.stackoverflow.co/2025/ai",
    )
    brief.core_properties.title = "Favorable Stance Is Not Accuracy Trust"
    brief.core_properties.subject = (
        "Executive brief on favorable stance, accuracy trust, and reported "
        "AI-use frequency"
    )
    brief.core_properties.author = "Josh Penzell"
    brief.core_properties.keywords = (
        "AI adoption, favorable stance, accuracy trust, current users"
    )
    brief.save(BRIEF_DOCX)


if __name__ == "__main__":
    required = [
        SENTIMENT_OUT / "current_user_daily_use_by_stance.csv",
        SENTIMENT_OUT / "full_sample_daily_use_by_stance.csv",
        SENTIMENT_OUT / "current_user_high_distrust_daily_use_by_stance.csv",
        SENTIMENT_OUT / "primary_current_user_model_comparison.csv",
        SENTIMENT_OUT / "full_sample_model_comparison.csv",
        SENTIMENT_OUT / "professional_current_user_model_comparison.csv",
        SENTIMENT_OUT / "primary_repeated_cv_comparisons.csv",
        SENTIMENT_OUT / "primary_country_grouped_cv.csv",
        SENTIMENT_OUT / "primary_heldout_metrics.json",
        SENTIMENT_OUT / "primary_current_user_stance_gradient.png",
        SENTIMENT_OUT / "primary_current_user_model_auc.png",
        SENTIMENT_OUT / "primary_current_user_stance_trust_heatmap.png",
        SENTIMENT_OUT / "three_year_context.png",
    ]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError(
            "Run analysis/sentiment_deep_dive.py first:\n"
            + "\n".join(missing)
        )
    build_report()
    build_brief()
    print(REPORT_DOCX)
    print(BRIEF_DOCX)
