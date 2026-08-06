from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PUBLICATION = ROOT / "publication"
OUTPUTS = ROOT / "outputs"
SENTIMENT = OUTPUTS / "sentiment"

BLOG_MD = PUBLICATION / "BLOG_POST.md"
BLOG_DOCX = PUBLICATION / "The_Developers_Using_AI_Without_Trusting_It_2026.docx"
NOTE_SOURCE = OUTPUTS / "Adoption_Without_Confidence_Research_Report_2026.docx"
NOTE_DOCX = OUTPUTS / "Adoption_Without_Confidence_Open_Research_Note_2026.docx"

INK = RGBColor(24, 43, 56)
BLUE = RGBColor(35, 91, 119)
DARK_BLUE = RGBColor(30, 70, 91)
GREEN = RGBColor(42, 125, 98)
MUTED = RGBColor(88, 96, 102)
LIGHT = RGBColor(232, 240, 238)
BLACK = RGBColor(0, 0, 0)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Imagination Applied  |  ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run2 = paragraph.add_run()
    run2._r.append(fld_char1)
    run2._r.append(instr_text)
    run2._r.append(fld_char2)
    set_run_font(run2, size=9, color=MUTED)


def add_hyperlink(paragraph, text, url, *, bold=False, italic=False):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "235B77")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if italic:
        r_pr.append(OxmlElement("w:i"))
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|\[[^\]]+\]\([^)]+\)|`[^`]+`)"
)


def add_inline(paragraph, text):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=11, color=INK)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=11, bold=True, color=INK)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=11, italic=True, color=INK)
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=10, color=DARK_BLUE, name="Courier New")
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=11, color=INK)


def create_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abstract = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abstract or [0]) + 1
    next_num = max(existing_num or [0]) + 1

    def add_definition(abstract_id, num_id, fmt, text):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "280")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "290")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Calibri")
        r_fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(r_fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), str(abstract_id))
        num.append(abstract_num_id)
        numbering.append(num)

    add_definition(next_abs, next_num, "bullet", "•")
    add_definition(next_abs + 1, next_num + 1, "decimal", "%1.")
    return next_num, next_num + 1


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("Figure Caption", "Source Note", "Pull Quote"):
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    caption = styles["Figure Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    source = styles["Source Note"]
    source.font.name = "Calibri"
    source.font.size = Pt(9)
    source.font.color.rgb = MUTED
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(4)

    pull = styles["Pull Quote"]
    pull.font.name = "Calibri"
    pull.font.size = Pt(13)
    pull.font.bold = True
    pull.font.color.rgb = DARK_BLUE
    pull.paragraph_format.left_indent = Inches(0.35)
    pull.paragraph_format.right_indent = Inches(0.35)
    pull.paragraph_format.space_before = Pt(10)
    pull.paragraph_format.space_after = Pt(10)
    pull.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_image(doc, path, alt_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.15))
    drawing = run._r.find(qn("w:drawing"))
    if drawing is not None:
        doc_pr = drawing.find(".//" + qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", alt_text)
            doc_pr.set("title", "Research figure")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def add_title_block(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(52)
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("IMAGINATION APPLIED | OPEN RESEARCH")
    set_run_font(run, size=10, bold=True, color=GREEN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(9)
    run = p.add_run(title)
    set_run_font(run, size=30, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(subtitle)
    set_run_font(run, size=15, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(26)
    run = p.add_run("Josh Penzell  |  Imagination Applied  |  July 2026")
    set_run_font(run, size=10.5, bold=True, color=MUTED)


def parse_front_matter(lines):
    if not lines or lines[0].strip() != "---":
        return {}, lines
    meta = {}
    i = 1
    while i < len(lines) and lines[i].strip() != "---":
        if ":" in lines[i]:
            key, value = lines[i].split(":", 1)
            meta[key.strip()] = value.strip().strip('"')
        i += 1
    return meta, lines[i + 1 :]


def build_blog_docx():
    raw_lines = BLOG_MD.read_text(encoding="utf-8").splitlines()
    meta, lines = parse_front_matter(raw_lines)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    configure_styles(doc)
    bullet_id, number_id = create_numbering(doc)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "THE DEVELOPERS USING AI WITHOUT TRUSTING IT"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(hp.runs[0], size=8.5, bold=True, color=MUTED)

    footer = section.footer
    set_page_number(footer.paragraphs[0])
    set_page_number(section.first_page_footer.paragraphs[0])

    title = meta.get("title", "The Developers Using AI Without Trusting It")
    subtitle = meta.get("subtitle", "")
    add_title_block(doc, title, subtitle)

    i = 0
    first_h1_skipped = False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# ") and not first_h1_skipped:
            first_h1_skipped = True
            i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1
            if i < len(lines) and lines[i].strip().startswith("*"):
                i += 1
            continue
        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p_pr = p._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "D6E2DF")
            p_bdr.append(bottom)
            p_pr.append(p_bdr)
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(stripped[4:], style="Heading 2")
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(stripped[3:], style="Heading 1")
            i += 1
            continue
        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            alt, rel_path = image_match.groups()
            add_image(doc, PUBLICATION / rel_path, alt)
            i += 1
            continue
        if stripped.startswith("> "):
            p = doc.add_paragraph(style="Pull Quote")
            add_inline(p, stripped[2:])
            i += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph()
            apply_numbering(p, bullet_id)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.208
            add_inline(p, stripped[2:])
            i += 1
            continue
        ordered = re.match(r"^\d+\.\s+(.*)", stripped)
        if ordered:
            p = doc.add_paragraph()
            apply_numbering(p, number_id)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.208
            add_inline(p, ordered.group(1))
            i += 1
            continue
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph(style="Figure Caption")
            p.add_run(stripped[1:-1])
            i += 1
            continue

        paragraph_lines = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].strip()
            if (
                not nxt
                or nxt.startswith("#")
                or nxt.startswith("- ")
                or re.match(r"^\d+\.\s+", nxt)
                or nxt.startswith(">")
                or nxt.startswith("![")
                or nxt == "---"
            ):
                break
            paragraph_lines.append(nxt)
            j += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_lines))
        i = j

    doc.core_properties.title = title
    doc.core_properties.subject = "Open research commentary on AI adoption, favorable stance, and accuracy trust"
    doc.core_properties.author = "Josh Penzell"
    doc.core_properties.keywords = "AI adoption; trust; favorable stance; Stack Overflow; developers"
    doc.core_properties.comments = "Publication-ready editorial copy, July 2026"
    doc.save(BLOG_DOCX)


def replace_text_in_paragraph(paragraph, replacements):
    full = "".join(run.text for run in paragraph.runs)
    revised = full
    for old, new in replacements:
        revised = revised.replace(old, new)
    if revised != full:
        if paragraph.runs:
            paragraph.runs[0].text = revised
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(revised)


def walk_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from walk_table(nested)


def repackage_research_note():
    doc = Document(NOTE_SOURCE)
    replacements = [
        ("RESEARCH REPORT", "OPEN RESEARCH NOTE"),
        ("Research Report", "Open Research Note"),
        ("IMAGINATION APPLIED RESEARCH SERIES", "IMAGINATION APPLIED OPEN RESEARCH SERIES"),
        ("Imagination Applied Research Series", "Imagination Applied Open Research Series"),
        ("This report uses", "This note uses"),
    ]
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for paragraph in walk_table(table):
            replace_text_in_paragraph(paragraph, replacements)
    for section in doc.sections:
        for part in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            for paragraph in part.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
            for table in part.tables:
                for paragraph in walk_table(table):
                    replace_text_in_paragraph(paragraph, replacements)
    doc.core_properties.title = "Adoption Without Confidence? Open Research Note"
    doc.core_properties.subject = (
        "Favorable stance, accuracy trust, and reported AI-use frequency "
        "in the 2025 Stack Overflow Developer Survey"
    )
    doc.core_properties.comments = "Open Research Note, version 3.1.0, July 2026"
    doc.save(NOTE_DOCX)


def copy_public_assets():
    assets = PUBLICATION / "assets"
    assets.mkdir(parents=True, exist_ok=True)
    shutil.copy2(
        SENTIMENT / "primary_current_user_stance_gradient.png",
        assets / "daily-use-by-stance.png",
    )
    shutil.copy2(
        SENTIMENT / "primary_current_user_model_auc.png",
        assets / "model-comparison.png",
    )


if __name__ == "__main__":
    copy_public_assets()
    build_blog_docx()
    repackage_research_note()
    print(BLOG_DOCX)
    print(NOTE_DOCX)
